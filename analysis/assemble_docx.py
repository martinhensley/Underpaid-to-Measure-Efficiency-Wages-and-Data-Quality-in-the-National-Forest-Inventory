#!/usr/bin/env python3
"""
Assemble the complete revised paper into a Word document (.docx).

Reads all revised section markdown files, embeds tables from CSV outputs,
inserts figures as images, and produces a formatted manuscript suitable
for journal submission.

Usage:
    python assemble_docx.py
"""

import csv
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

BASE = Path(__file__).parent.parent
TABLE_DIR = BASE / "tables"
FIG_DIR = BASE / "figures"
OUT_PATH = BASE / "Underpaid_to_Measure_revised.docx"


def set_style(doc: Document):
    """Configure document-level styles."""
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(12)
    pf = style.paragraph_format
    pf.space_after = Pt(6)
    pf.line_spacing = 2.0  # double-spaced for submission


def add_title_page(doc: Document):
    doc.add_paragraph()  # spacer
    doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Underpaid to Measure? Efficiency Wages and\nData Quality in the National Forest Inventory")
    run.bold = True
    run.font.size = Pt(16)

    doc.add_paragraph()

    author = doc.add_paragraph()
    author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = author.add_run("Martin Hensley")
    run.font.size = Pt(14)

    affil = doc.add_paragraph()
    affil.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = affil.add_run("Independent Researcher")
    run.font.size = Pt(12)
    run.italic = True

    doc.add_paragraph()

    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_p.add_run("March 2026")
    run.font.size = Pt(12)

    doc.add_paragraph()
    doc.add_paragraph()

    # Keywords and JEL Classification
    kw = doc.add_paragraph()
    kw.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = kw.add_run("Keywords: ")
    run.bold = True
    kw.add_run("efficiency wages, moral hazard, forest inventory, data quality, ocular estimation, principal-agent, FIA")

    jel = doc.add_paragraph()
    jel.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = jel.add_run("JEL Classification: ")
    run.bold = True
    jel.add_run("J41, Q23, D82, D83, Q28")

    doc.add_page_break()


def add_abstract(doc: Document):
    h = doc.add_heading("Abstract", level=1)
    h.runs[0].font.size = Pt(14)

    doc.add_paragraph(
        "Using paired QA/production measurements of 94,459 trees across 24 "
        "northern states (Yanai et al., 2023), we find that production field "
        "crews in the USDA Forest Inventory and Analysis (FIA) program record "
        "ocular estimation for 23.1% of tree heights versus 5.8% for QA "
        "crews \u2014 a 4:1 ratio providing direct evidence of differential "
        "measurement effort under different institutional conditions. Height "
        "discrepancies between crew types are dramatic (mean |diff| = 3.31 ft) "
        "while diameter discrepancies are minimal (0.055 in.), consistent with "
        "effort reduction concentrating on the hard-to-verify margin as "
        "predicted by multi-task moral hazard theory (Holmstr\u00f6m and Milgrom, "
        "1991). The estimation differential is pervasive year-round rather "
        "than concentrated in late-season months \u2014 consistent with the finite-horizon prediction that terminal seasonal workers face violated incentive constraints throughout their appointment, not only at season\u2019s end."
    )

    doc.add_paragraph(
        "A supporting efficiency wage calibration (Shapiro and Stiglitz, "
        "1984) using federal pay schedules, stated quality assurance audit "
        "rates, and competing labor market wages shows that the no-shirking "
        "condition (NSC) substantially exceeds current technician wages "
        "under central parameter assumptions, though the magnitude of the "
        "gap is sensitive to the effort cost parameter, which has the "
        "weakest direct empirical grounding. The calibrated model identifies "
        "low effective detection probability as the parameter most strongly "
        "governing the incentive gap and shows that increasing effective "
        "detection probability through expanded audits and "
        "technology-assisted monitoring is substantially more cost-effective "
        "than raising wages within the General Schedule structure."
    )

    doc.add_page_break()


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h


def add_section1(doc: Document):
    """Introduction."""
    add_heading(doc, "1. Introduction")

    paras = [
        "The United States Forest Inventory and Analysis (FIA) program is the "
        "nation\u2019s continuous forest census \u2014 a systematic inventory of trees "
        "across all ownerships that supports approximately $80 billion in annual "
        "timber commerce, informs national carbon accounting, and underpins forest "
        "management decisions at every scale from timber sales to climate policy. "
        "The program employs several hundred seasonal field technicians each year "
        "to visit a nationally representative sample of approximately 130,000 "
        "permanent plots on a rotating five- to ten-year cycle, measuring individual "
        "tree diameters, heights, species, and condition. The integrity of this "
        "information infrastructure depends almost entirely on the accuracy of "
        "these field measurements.",

        "This paper examines whether the institutional design of the FIA program "
        "creates conditions under which rational field technicians would "
        "systematically approximate measurements \u2014 substituting ocular "
        "estimates for instrument-based measurements. Using paired QA/production "
        "data from Yanai et al. (2023) \u2014 94,459 trees across 24 states "
        "measured independently by both production and QA crews \u2014 we find "
        "that production crews self-report ocular estimation (HTCD=3) for "
        "23.1% of heights, compared to 5.8% for QA crews. This 4:1 ratio, "
        "recorded in literal method codes rather than inferred from statistical "
        "forensics, provides direct evidence that measurement effort differs "
        "systematically across institutional conditions.",

        "We make three contributions. First, we document differential estimation "
        "across crew types using paired QA/production data, providing the most "
        "direct evidence to date that FIA production crews use ocular estimation "
        "at substantially elevated rates. The height-diameter asymmetry in "
        "measurement discrepancies \u2014 dramatic for heights, minimal for "
        "diameters \u2014 confirms that effort reduction concentrates on the "
        "hard-to-verify margin, as multi-task moral hazard theory predicts "
        "(Holmstr\u00f6m and Milgrom, 1991).",

        "Second, we develop a principal-agent model drawing on the efficiency "
        "wage framework of Shapiro and Stiglitz (1984) as a motivating benchmark "
        "for interpreting these findings. The calibrated no-shirking condition "
        "(NSC) identifies the institutional parameters governing incentive "
        "alignment \u2014 particularly the low effective detection probability "
        "\u2014 and provides quantitative targets for policy intervention.",

        "Third, we derive specific policy recommendations from the calibrated "
        "model: monitoring intensities, institutional reforms, and technology-"
        "assisted verification approaches that would narrow the gap between "
        "production and QA measurement conditions.",

        "Our analysis connects to the economics of monitoring and incentives "
        "(Shapiro and Stiglitz, 1984; Holmstr\u00f6m, 1979; Holmstr\u00f6m and "
        "Milgrom, 1991), the intrinsic motivation literature (B\u00e9nabou and "
        "Tirole, 2003, 2006), and natural resource monitoring, where Duflo et "
        "al. (2013) demonstrate experimentally that auditor incentive design "
        "dramatically affects environmental compliance reporting in India. "
        "To our knowledge, this paper represents a novel application of the "
        "Shapiro-Stiglitz efficiency wage framework specifically to natural "
        "resource data collection programs, complementing the Duflo et al. "
        "focus on third-party auditing.",

        "The remainder of the paper is organized as follows. Section 2 provides "
        "institutional background and develops the motivating theoretical "
        "framework. Section 3 presents the empirical analysis. Section 4 "
        "discusses implications for timber markets and data governance. "
        "Section 5 derives policy recommendations. Section 6 discusses "
        "limitations, and Section 7 concludes.",
    ]
    for p in paras:
        doc.add_paragraph(p)


def add_section2(doc: Document):
    """Institutional Background and Theoretical Framework."""
    add_heading(doc, "2. Institutional Background and Theoretical Framework")

    add_heading(doc, "2.1 The FIA Program", level=2)
    paras = [
        "The Forest Inventory and Analysis program, administered by the USDA "
        "Forest Service, conducts the national forest inventory mandated by the "
        "Forest and Rangeland Renewable Resources Research Act of 1978 (16 U.S.C. "
        "\u00a71641\u20131646). Since the transition from periodic to annual inventory "
        "in the early 2000s, FIA has operated a panel-based sampling design in "
        "which approximately one-fifth to one-seventh of permanent plot locations "
        "are visited each year. The national plot network consists of approximately "
        "130,000 permanent sample plots on a hexagonal grid with one plot per "
        "approximately 6,000 acres of forest land (Bechtold and Patterson, 2005).",

        "Each plot consists of four circular subplots (24-foot radius for trees "
        "\u22655.0 inches DBH, with nested microplots for seedlings and saplings). "
        "At each subplot, field crews measure all qualifying trees: diameter at "
        "breast height (DBH) with a diameter tape to the nearest 0.1 inch, total "
        "height with a clinometer or laser hypsometer to the nearest foot, species "
        "identification, crown class and condition, and numerous auxiliary "
        "variables. A single production plot typically requires 4\u20138 hours to "
        "complete, including travel, navigation, and data entry.",

        "The measurement season runs roughly April through November in most "
        "regions. Southern states may begin earlier; northern and mountain states "
        "may extend into early winter depending on snow conditions. This seasonal "
        "window creates a natural cycle of fieldwork intensity central to our "
        "analysis.",
    ]
    for p in paras:
        doc.add_paragraph(p)

    add_heading(doc, "2.2 Quality Assurance Structure", level=2)
    paras = [
        "FIA operates a two-tier quality assurance program. Hot checks are "
        "supervisory observations where a crew leader accompanies a field "
        "technician and evaluates adherence to measurement protocols. These are "
        "relatively frequent but not independent \u2014 the technician knows they are "
        "being observed. Blind checks are independent re-measurements where a "
        "separate QA crew revisits a randomly selected production plot without the "
        "original crew\u2019s knowledge. Discrepancies between production and QA "
        "measurements are compared to tolerance thresholds defined in the FIA "
        "National Core Field Guide. The stated target is that approximately 4% of "
        "production plots receive blind-check QA re-measurement each year (USDA "
        "Forest Service, 2023).",

        "From the technician\u2019s perspective, the effective detection probability "
        "is the product of (1) the probability that the specific plot is selected "
        "for QA, (2) the probability that the QA re-measurement is conducted "
        "before the technician\u2019s appointment ends, and (3) the probability that "
        "the comparison detects approximation conditional on it occurring. Even at "
        "the stated 4% annual rate, this compound probability is substantially "
        "less than 0.04 per plot-period.",
    ]
    for p in paras:
        doc.add_paragraph(p)

    add_heading(doc, "2.3 The Seasonal Labor Market", level=2)
    paras = [
        "FIA field crews are predominantly seasonal (term) federal employees "
        "hired under the General Schedule at GS-4 (entry) through GS-6 "
        "(experienced crew leader) pay grades. Term appointments under 5 CFR "
        "\u00a7316.401 are limited to 13 months and carry no guarantee of rehire. "
        "In practice, successful technicians are often rehired across multiple "
        "seasons, but the employment relationship is explicitly temporary.",

        "This seasonal structure creates three features relevant to incentive "
        "analysis. First, the effective outside option includes the full set of "
        "seasonal outdoor employment opportunities \u2014 wildfire suppression crews "
        "(often at higher base pay plus hazard pay), trail maintenance, seasonal "
        "NPS and BLM positions, and state forestry agency work. BLS Occupational "
        "Employment and Wage Statistics reveal substantial geographic variation: "
        "median hourly wages for landscaping range from approximately $15/hr in "
        "the Southeast to $19/hr in the Pacific Northwest, while logging worker "
        "wages range from $17/hr to $22/hr.",

        "Second, the guaranteed end of the seasonal appointment creates an "
        "end-of-season dynamic that undermines the efficiency wage mechanism. As "
        "the appointment end date approaches, the termination threat loses potency "
        "\u2014 the technician is losing the job regardless. The effective separation "
        "rate b increases toward the end of the season, raising the NSC precisely "
        "when cumulative fatigue and declining weather may also be reducing "
        "intrinsic motivation.",

        "Third, mid-season termination for performance reasons, while legally "
        "permissible, faces practical barriers: documentation requirements, "
        "progressive discipline expectations, and the difficulty of demonstrating "
        "that approximated data constitutes inadequate performance when the data "
        "superficially meets plausibility standards.",
    ]
    for p in paras:
        doc.add_paragraph(p)

    # 2.4 A Motivating Model
    add_heading(doc, "2.4 A Motivating Model", level=2)
    paras = [
        "We use the Shapiro-Stiglitz (1984) efficiency wage framework as a "
        "quantitative lens for assessing whether institutional design creates "
        "conditions favoring measurement approximation \u2014 not a causal claim "
        "about why the HTCD differential documented in Section 3 exists. The "
        "calibration provides an indicative benchmark for the magnitude of a "
        "potential incentive gap, not the precise behavioral mechanism through "
        "which that gap translates into measurement decisions.",

        "Setup. A risk-neutral field technician employed at wage w per period "
        "chooses effort e \u2208 {0, \u0113}: careful measurement (e = \u0113, "
        "disutility \u0113) or approximation (e = 0). The principal detects "
        "approximation with probability q per period through QA audits. Upon "
        "detection, the technician is terminated and earns outside option \u0175. "
        "With exogenous separation rate b and discount rate r, the standard "
        "value function analysis yields the no-shirking condition (NSC). "
        "Intuitively, if workers can cut corners without being caught, they "
        "will only refrain if the job pays enough above their next-best option "
        "to make the risk of termination costly. The minimum wage that makes "
        "honest effort worthwhile is:",

        "    w* \u2265 \u0175 + \u0113 + (\u0113/q)(r + b)",

        "The critical wage has three components: the outside option \u0175; a "
        "compensating differential for effort \u0113; and the efficiency wage "
        "premium (\u0113/q)(r + b) \u2014 the extra pay needed to make the job "
        "valuable enough that workers fear losing it. This premium is large "
        "when detection is infrequent (q small), the job is ending soon "
        "(b large), or the discount rate is high. Risk aversion would increase "
        "the NSC, so our risk-neutral specification is conservative.",

        "Finite-horizon extension. The stationary SS model assumes a constant "
        "hazard of termination, whereas seasonal FIA employment has a known "
        "endpoint. In a finite-horizon setting, the termination threat is "
        "vacuous in the final period \u2014 the worker is leaving regardless "
        "\u2014 and by backward induction, the NSC is violated in every "
        "period. The standard resolution in the repeated-game literature is "
        "that reputation or continuation value sustains cooperation, but "
        "these mechanisms require that the employment relationship extends "
        "beyond the current season. For a purely terminal seasonal worker "
        "with no prospect of rehire, backward induction implies that "
        "estimation is rational throughout the appointment, not merely at "
        "season\u2019s end. This is actually more consistent with the "
        "empirical finding of year-round estimation (Section 3.2) than the "
        "stationary model, which would predict late-season concentration as "
        "the effective separation rate b increases. We therefore interpret "
        "the calibration as a benchmark indicating the direction and "
        "magnitude of incentive gaps, not as a precise wage threshold.",

        "For career-track workers \u2014 GS-0462 technicians pursuing "
        "permanent GS-0460 positions through the "
        "GS-4\u2192GS-5\u2192GS-7 pipeline \u2014 the employment "
        "relationship is effectively indefinite-horizon with periodic "
        "re-contracting. The substantial employment rents of permanent "
        "federal service (FERS retirement, FEHB health insurance, job "
        "security) create continuation value that sustains cooperation "
        "across seasons. For this subpopulation, the stationary SS "
        "framework is a reasonable approximation: the multi-season career "
        "structure means the repeated-game logic applies over a multi-year "
        "horizon. The calibrated NSC therefore applies most directly to "
        "terminal seasonals without career pathways \u2014 the vulnerable "
        "margin where neither backward induction nor career incentives "
        "sustain measurement effort.",
    ]
    for p in paras:
        doc.add_paragraph(p)

    # Table 1: Parameters
    add_csv_table(doc, TABLE_DIR / "table1_parameters.csv",
                  "Table 1: Model Parameters and Sources")

    paras = [
        "Wages follow the 2025 General Schedule with Rest of U.S. (RUS) "
        "locality pay (OPM, 2025). The outside option (\u0175 = $16/hr) "
        "reflects a central estimate across competing seasonal occupations. "
        "The effort cost (\u0113 = 0.25 \u00d7 $17.27 = $4.32/hr) is grounded "
        "in a time-budget micro-foundation: instrument height measurement "
        "requires approximately 2\u20134 minutes per tree versus 15\u201330 seconds "
        "for ocular estimation; across ~35 height-measured trees per plot, "
        "substituting estimation saves 60\u2013175 minutes per 4\u20138 hour "
        "plot day.",

        "The detection probability (q = 0.01 per month) is deliberately "
        "conservative \u2014 it overstates the effective detection probability "
        "for most states. We computed effective state-level QA rates as the "
        "ratio of blind-check plots (QA_STATUS = 7) to production plots "
        "(QA_STATUS = 1) in the FIA DataMart PLOT table, computed separately "
        "by state and measurement year and averaged across the sample period. "
        "State-level rates range from 0.0001 (VT) to 0.013 (GA, OR), with a "
        "cross-state median of approximately 0.003. The order-of-magnitude "
        "gap between the observed median (~0.3%) and the stated 4% target "
        "reflects both incomplete coverage and the compound probability "
        "structure described above. At this observed median, the NSC rises to "
        "approximately $54.92/hr.",

        "NSC evaluation. At central parameters:\n"
        "    w* = 16 + 4.32 + (4.32/0.01)(0.004 + 0.02) = $30.69/hr",

        "This benchmark value is 78% above the GS-4 wage ($17.27/hr), 59% "
        "above the GS-5 wage ($19.33/hr), and 43% above the GS-6 wage "
        "($21.54/hr). The efficiency wage premium ($10.37/hr) dominates, "
        "driven by low detection probability. If q increases to 0.04 (the "
        "stated FIA QA rate applied monthly), w* falls to approximately "
        "$22.91/hr, within reach of GS-7 ($23.94/hr) wages. However, "
        "q = 0.04 represents the stated FIA target rate; the observed "
        "effective rate from FIA DataMart is approximately 0.003 across "
        "most states, so this favorable scenario requires closing the "
        "order-of-magnitude gap between stated and effective monitoring "
        "rates. This sensitivity underscores that the detection probability "
        "q is the parameter most strongly governing the incentive gap.",
    ]
    for p in paras:
        doc.add_paragraph(p)

    doc.add_paragraph(
        "The magnitude of the NSC gap is also sensitive to the effort cost "
        "assumption: at \u0113 = $2/hr (the lower bound of plausible values), "
        "w* falls to approximately $22/hr, within reach of GS-6. Table 2b "
        "presents the full sensitivity surface. The qualitative conclusion "
        "\u2014 that the NSC exceeds current wages under central and most "
        "alternative parameterizations \u2014 is robust, but specific "
        "percentage gaps should be interpreted as indicative benchmarks, "
        "not precise thresholds."
    )

    # Table 2: Sensitivity
    add_csv_table(doc, TABLE_DIR / "table2_nsc_sensitivity.csv",
                  "Table 2: NSC Sensitivity to Outside Option (\u0175) and Separation Rate (b)")

    # Table 2b: q sensitivity
    add_csv_table(doc, TABLE_DIR / "table2b_nsc_q_sensitivity.csv",
                  "Table 2b: NSC Sensitivity to Detection Probability (q) and Effort Cost (\u0113)")

    # Table 3: Cross-state
    add_csv_table(doc, TABLE_DIR / "table3_cross_state_nsc.csv",
                  "Table 3: Cross-State NSC Variation Using BLS Wage Data")

    # Figure 1
    add_figure(doc, FIG_DIR / "figure1_nsc_sensitivity.png",
               "Figure 1: NSC Sensitivity to Detection Probability (q) and Effort Cost")

    # 2.5 Multi-Task Predictions and Behavioral Extensions
    add_heading(doc, "2.5 Multi-Task Predictions and Behavioral Extensions", level=2)
    paras = [
        "Multi-task effort allocation. When workers perform multiple tasks "
        "that differ in how easily supervisors can verify them, theory predicts "
        "effort will decline most on the hardest-to-monitor tasks. The binary "
        "effort model treats \u201ccareful measurement\u201d as a single "
        "dimension. In practice, FIA "
        "technicians allocate effort across multiple measurement tasks with "
        "different costs and detection probabilities \u2014 a multi-task moral "
        "hazard problem (Holmstr\u00f6m and Milgrom, 1991). DBH measurement "
        "with a diameter tape is easily verified by QA re-taping; height "
        "measurement with a clinometer requires retreating to find line-of-sight "
        "and identifying the correct top in closed canopy \u2014 and has "
        "inherently higher inter-observer variance. Species identification "
        "and condition codes involve partly subjective categorical judgments. "
        "The multi-task framework predicts that effort reductions concentrate "
        "on the hardest-to-verify margins \u2014 particularly height estimation. "
        "This prediction is directly testable with the paired data in "
        "Section 3.",

        "Intrinsic motivation. Many FIA technicians are trained foresters with "
        "genuine professional commitment (Perry and Wise, 1990; Buurman et al., "
        "2012). Following B\u00e9nabou and Tirole (2003, 2006), intrinsic "
        "motivation m effectively reduces the net effort cost to (\u0113 \u2212 m); "
        "when m \u2265 \u0113, the NSC is automatically satisfied. However, "
        "intrinsic motivation faces structural erosion: seasonal employment at "
        "GS-4/5 wages undermines professional identity and organizational "
        "attachment; cumulative fatigue reduces the psychic rewards of careful "
        "measurement; and crowding out (B\u00e9nabou and Tirole, 2003) means "
        "monitoring perceived as distrustful can itself reduce m. Our policy "
        "recommendation is not increased surveillance beyond stated expectations "
        "but credible monitoring at the existing stated rate.",

        "Workforce heterogeneity. The FIA seasonal workforce contains at "
        "least two distinct types. Career-track seasonals \u2014 GS-0462 "
        "technicians pursuing permanent GS-0460 positions \u2014 treat the "
        "seasonal appointment as a pipeline into federal careers with "
        "substantial employment rents (FERS retirement, FEHB health insurance). "
        "The multi-season career structure (GS-4\u2192GS-5 after one season, "
        "GS-7 upon degree completion) means the SS framework\u2019s repeated-game "
        "logic applies over a multi-year horizon. Terminal seasonals, by "
        "contrast, lack pathways to permanent employment and face the model\u2019s "
        "standard NSC. The calibrated NSC applies to the marginal (terminal) "
        "worker. Career-track workers\u2019 tendency to stay late in the season "
        "means late-season crews are positively selected on effort quality, "
        "attenuating seasonal signals.",
    ]
    for p in paras:
        doc.add_paragraph(p)


def add_section3(doc: Document):
    """Empirical Analysis (renumbered from old Section 4)."""
    add_heading(doc, "3. Empirical Analysis")

    doc.add_paragraph(
        "The motivating model in Section 2 identifies institutional conditions "
        "under which the no-shirking condition is violated. This section tests "
        "the model\u2019s predictions using paired QA/production measurements "
        "from Yanai et al. (2023), which provide direct evidence of differential "
        "measurement effort across crew types operating under different "
        "institutional incentives. Supplementary forensic analyses \u2014 digit "
        "heaping, allometric residual patterns, and remeasurement growth "
        "anomalies \u2014 largely return null seasonal results, consistent with "
        "the finding below that estimation is pervasive year-round rather than "
        "seasonally concentrated; these are reported in the Supplementary "
        "Materials (Sections A.1\u2013A.7)."
    )

    # 3.1 Data
    add_heading(doc, "3.1 Data", level=2)
    paras = [
        "We use the paired QA/production dataset from Yanai et al. (2023, "
        "doi:10.2737/RDS-2022-0056), which provides 94,459 paired tree "
        "observations across 24 northern states (2011\u20132016). Each tree was "
        "measured independently by both a production (field) crew and a QA crew, "
        "with measurements recorded under prefixed field names (F_ for production, "
        "Q_ for QA). The dataset includes HTCD (height method code), recording "
        "whether each height was measured with instruments (HTCD=1), modeled "
        "(HTCD=2), or visually estimated (HTCD=3).",

        "Restricting to live trees present in both visits yields 37,717 paired "
        "diameter measurements and 31,568 paired height measurements with valid "
        "HTCD codes for both crews. The geographic distribution is broad: 24 "
        "states spanning the Northern, Northeastern, and North Central FIA "
        "regions, with no single state contributing more than 12% of paired "
        "observations. Critically, the dataset does not include Georgia \u2014 "
        "eliminating the geographic concentration that limits inference in some "
        "supplementary analyses.",

        "Production and QA crews operate under fundamentally different "
        "institutional conditions. Production crews are predominantly seasonal "
        "(term) GS-4/5 employees working under throughput pressure. QA crews "
        "are typically senior technicians or supervisory staff with no production "
        "quotas, no seasonal appointment pressure, and an explicit quality "
        "mandate. This institutional contrast provides the identifying variation.",

        "We define LATE_SEASON as an indicator for months 9\u201312. Regressions "
        "use OLS with standard errors clustered at the state level. Additional "
        "forensic analyses using FIA DataMart data from eight states (4.6 million "
        "tree-observations) are reported in the Supplementary Materials.",
    ]
    for p in paras:
        doc.add_paragraph(p)

    # 3.2 Height Method Codes
    add_heading(doc, "3.2 Height Method Codes: Direct Evidence of Differential Estimation", level=2)
    paras = [
        "The headline finding is the differential rate of visual estimation. "
        "Production crews record HTCD=3 (ocular estimation) for 23.1% of heights; "
        "QA crews record HTCD=3 for only 5.8% \u2014 a 4:1 ratio. This gap is "
        "present across species types: for conifers, the production HTCD=3 rate "
        "is 25.9% versus 5.1% for QA; for hardwoods, 22.1% versus 6.0%.",

        "The cross-tabulation of F_HTCD \u00d7 Q_HTCD reveals a striking "
        "asymmetry. Of the paired observations, 6,402 trees were estimated by "
        "production crews but measured by QA crews, while only 1,010 were "
        "measured by production but estimated by QA \u2014 a 6.3:1 asymmetry "
        "in the direction predicted by the model.",

        "Not all ocular estimation is unauthorized. FIA protocols permit visual "
        "estimation in specific circumstances. If both production and QA crews "
        "face similar terrain, the differential in HTCD=3 rates (23.1% vs 5.8%) "
        "provides a lower bound on the excess estimation attributable to "
        "institutional differences rather than field conditions.",

        "This differential is not seasonal. A linear probability model of "
        "Prob(F_HTCD=3) on LATE_SEASON with state and year fixed effects yields "
        "a coefficient of \u22120.015 (SE = 0.024, p = 0.54). Production crews "
        "use ocular estimation at roughly the same elevated rate throughout the "
        "measurement season. Neither conifers (p = 0.22) nor hardwoods "
        "(p = 0.58) show a significant seasonal pattern.",

        "The absence of seasonal variation is consistent with structural incentive "
        "misalignment \u2014 if the NSC is violated at current wages throughout the "
        "season, estimation should be pervasive rather than concentrated in late "
        "months. The evidentiary weight rests on the 4:1 production-QA "
        "differential, not on the seasonal pattern.",

        "State-level variation in HTCD=3 rates provides additional context. While "
        "all states in the sample show higher production than QA estimation rates, "
        "the magnitude varies \u2014 suggesting that local institutional factors "
        "(crew training, supervisory culture, terrain difficulty) modulate the "
        "overall incentive structure. In principle, cross-state variation in "
        "HTCD=3 rates could provide a partial mechanism test \u2014 whether the "
        "estimation differential is larger in states with wider calibrated NSC "
        "gaps (Table 3). However, with only 24 states and strong confounding "
        "between outside wages, terrain difficulty, and forest type, the "
        "statistical power for such a test is limited; we therefore note the "
        "pattern without drawing inferential conclusions.",
    ]
    for p in paras:
        doc.add_paragraph(p)

    # Paired QA/Production Summary Table
    t_paired = doc.add_table(rows=5, cols=4)
    t_paired.style = "Table Grid"
    t_paired.alignment = WD_TABLE_ALIGNMENT.CENTER
    t_paired_headers = ["Measure", "Production", "QA", "Ratio"]
    for i, h in enumerate(t_paired_headers):
        t_paired.rows[0].cells[i].text = h
        for p in t_paired.rows[0].cells[i].paragraphs:
            for r in p.runs:
                r.bold = True
    t_paired_data = [
        ["HTCD = 3 rate", "23.1%", "5.8%", "4.0:1"],
        ["Mean |HT discrepancy| (ft)", "3.31", "\u2014", "\u2014"],
        ["Mean |DBH discrepancy| (in)", "0.06", "\u2014", "\u2014"],
        ["Seasonal variation in HTCD=3", "p = 0.54", "p = 0.54", "\u2014"],
    ]
    for row_idx, row_data in enumerate(t_paired_data, start=1):
        for col_idx, val in enumerate(row_data):
            t_paired.rows[row_idx].cells[col_idx].text = val

    cap_paired = doc.add_paragraph()
    run = cap_paired.add_run(
        "Table 4: Paired QA/Production Comparison Summary. "
        "Yanai et al. (2023) dataset, 24 northern states, 2011\u20132016."
    )
    run.bold = True
    run.italic = True
    cap_paired.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    add_figure(doc, FIG_DIR / "paired_qa_htcd_by_month.png",
               "Figure 2: HTCD = 3 (Ocular Estimation) Rates by Month: "
               "Production vs. QA Crews. Yanai et al. (2023) paired data, "
               "24 northern states, 2011\u20132016.")

    # 3.3 Corroborating Evidence
    add_heading(doc, "3.3 Corroborating Evidence: The Height-Diameter Asymmetry", level=2)
    paras = [
        "The multi-task model (Section 2.5) predicts that effort reduction "
        "concentrates on the hard-to-verify margin \u2014 height rather than "
        "diameter. The paired data confirm this dramatically.",

        "Measurement discrepancies. Height discrepancies between production and "
        "QA measurements are substantial: the mean absolute difference is 3.31 "
        "feet (SD = 4.91), with a slight negative bias (production heights "
        "average 0.15 feet below QA heights). The distribution is right-skewed, "
        "with a median absolute difference of 2.0 feet and a 90th percentile "
        "of 7.0 feet. For diameter, the mean absolute "
        "difference is only 0.055 inches (SD = 0.21) \u2014 confirming that the "
        "height margin is dramatically noisier. Neither "
        "discrepancy shows a significant seasonal pattern (height: p = 0.64; "
        "diameter: p = 0.76).",

        "Accuracy by method. Trees where production crews used instruments "
        "(HTCD=1) have mean |HT_diff| = 3.23 feet, while trees where production "
        "crews estimated (HTCD=3) have mean |HT_diff| = 3.48 feet \u2014 estimated "
        "heights are slightly less accurate, though the difference is modest.",

        "Allometric conformity. We fit an allometric model "
        "(ln(Q_HT) ~ ln(Q_DIA) + C(species group)) on QA measurements as "
        "the gold standard (R\u00b2 = 0.71). Trees with HTCD=1 (measured) "
        "have mean |residual| = 0.164 and trees with HTCD=3 (estimated) "
        "have mean |residual| = 0.166 \u2014 effectively identical "
        "(t = \u22120.77, p = 0.44). Estimated and measured heights show "
        "similar deviations from the species-specific height-diameter "
        "relationship, indicating that experienced technicians who estimate "
        "heights do so with sufficient skill that their estimates are "
        "forensically indistinguishable from instrument measurements via "
        "allometric residual analysis. This null is informative for "
        "monitoring design: it means that allometric outlier detection "
        "\u2014 flagging trees whose heights deviate from the expected "
        "diameter-height curve \u2014 would not identify ocular estimation, "
        "because skilled estimators produce values that conform to the "
        "population allometric relationship.",

        "The allometric conformity null also raises a welfare question: if "
        "estimated heights are nearly as accurate as instrument "
        "measurements, the data quality cost of estimation may be more "
        "nuanced than \u201capproximation equals error.\u201d To quantify: "
        "estimated trees (HTCD=3) have mean absolute height discrepancy "
        "of 3.48 ft versus 3.23 ft for measured trees \u2014 a modest 7.7% "
        "accuracy penalty per tree. When propagated through standard volume "
        "equations (volume \u221d height \u00d7 basal area), this translates "
        "to approximately 3\u20134% volume estimation error per tree. While "
        "individually modest, systematic estimation across 23% of "
        "production heights compounds across the inventory and introduces "
        "non-random error correlated with crew institutional conditions. "
        "Three considerations further temper an optimistic reading. First, "
        "the conformity "
        "null is a population average \u2014 the tails of the residual "
        "distribution may differ for site-specific applications (e.g., "
        "timber sales on unusual sites where allometric norms are poor "
        "predictors), even if the central tendency is well-preserved. "
        "Second, even if individual estimates are adequate on average, "
        "undetectable estimation erodes the monitoring architecture: when "
        "the QA system cannot distinguish estimated from measured heights, "
        "the equilibrium estimation rate can only increase over time as "
        "workers learn that non-compliance carries no consequences \u2014 "
        "creating a moral hazard spiral. Third, protocol compliance has "
        "independent value because downstream users rely on HTCD codes to "
        "assess data quality; if estimation is forensically invisible, the "
        "HTCD=1 code ceases to be an informative signal.",

        "More fundamentally, if skilled estimation is forensically "
        "undetectable through output-based monitoring \u2014 as the "
        "allometric conformity null suggests \u2014 the quality assurance "
        "problem shifts from \u201cdetect bad data\u201d to \u201cverify "
        "process compliance.\u201d This requires input monitoring (GPS track "
        "logging, time-on-plot verification) rather than output monitoring "
        "(allometric outlier detection, QA remeasurement comparison). The "
        "institutional corrosion concern is therefore not that estimated "
        "data are necessarily inaccurate, but that an output-based QA "
        "system is structurally incapable of maintaining protocol "
        "compliance when skilled estimation produces statistically "
        "indistinguishable results. We return to the policy implications "
        "of this distinction in Section 4.",

        "Remeasurement anomalies. Corroborating evidence from FIA\u2019s panel "
        "structure confirms the height-diameter asymmetry at scale: across "
        "2,464,859 remeasurement pairs in eight states, the height anomaly rate "
        "is 9.3% compared to 1.5% for DBH \u2014 a 6:1 ratio. The cleanest "
        "seasonal test \u2014 conifers only, eliminating the leaf-off visibility "
        "confound \u2014 yields a LATE_SEASON coefficient indistinguishable from "
        "zero (p = 0.98).",
    ]
    for p in paras:
        doc.add_paragraph(p)

    add_figure(doc, FIG_DIR / "paired_qa_discrepancies.png",
               "Figure 3: Production\u2013QA Measurement Discrepancies. "
               "(a) Height discrepancy (feet); (b) DBH discrepancy (inches).")

    # 3.4 Discussion
    add_heading(doc, "3.4 Discussion", level=2)
    paras = [
        "The empirical findings can be stated concisely: production crews "
        "estimate heights at four times the QA rate (23.1% vs. 5.8%), this "
        "differential is pervasive year-round, and effort reduction "
        "concentrates on height rather than diameter \u2014 the measurement "
        "dimension that is both hardest to verify and most time-costly to "
        "perform carefully.",

        "Multiple mechanisms are consistent with these patterns. The "
        "production-QA differential may reflect incentive misalignment "
        "(the NSC framework in Section 2.4), selection and experience "
        "differences (QA crews are typically GS-7+ supervisory staff), "
        "differential training emphasis, task framing (QA crews know "
        "measurements will be compared), or workload pressure from "
        "production quotas. Clinometer height measurement requires "
        "identifying a sightline, retreating to distance, taking dual "
        "angle readings, and computing \u2014 approximately 2\u20134 minutes "
        "per tree versus 30\u201360 seconds for DBH taping, a roughly 3:1 "
        "time-cost ratio (consistent with the \u0113 micro-foundation in "
        "Section 2.4). A crew under pure throughput pressure would "
        "therefore rationally concentrate time savings on height "
        "independent of any detection-probability calculus. The "
        "height-diameter asymmetry is consistent with both strategic "
        "effort allocation and time-pressure optimization; we cannot "
        "discriminate between these mechanisms from the measurement "
        "data alone.",

        "The finite-horizon analysis (Section 2.4) provides a "
        "complementary and possibly more fundamental explanation for the "
        "year-round estimation pattern. For terminal seasonal workers "
        "with no prospect of rehire, backward induction implies the NSC "
        "is violated in every period regardless of wage levels \u2014 "
        "estimation is rational throughout the appointment. This "
        "parameter-free prediction is directly confirmed by the absence "
        "of seasonal variation in HTCD=3 rates, and does not depend on "
        "the effort cost calibration.",

        "Whether the moral hazard spiral discussed in Section 3.3 has "
        "manifested empirically \u2014 i.e., whether HTCD=3 rates have "
        "trended upward over time \u2014 is an important question that "
        "the current data cannot definitively answer. The paired dataset "
        "covers only six years (2011\u20132016), too short to identify a "
        "slow-moving institutional drift. The FIA DataMart contains HTCD "
        "records extending back to approximately 1999, and was used in this "
        "study for supplementary forensic analyses (Sections A.1\u2013A.7). A "
        "systematic temporal trend analysis of HTCD=3 rates over this longer "
        "series was beyond the scope of the current study but represents a "
        "priority for future research.",

        "The policy prescription holds regardless of mechanism: whether "
        "production crews estimate more because of weaker incentives, "
        "less experience, time pressure, or some combination, the "
        "response is institutional redesign that narrows the gap between "
        "production and QA measurement conditions \u2014 through enhanced "
        "monitoring, reduced workload pressure, improved training, or "
        "career-ladder investments.",

        "The corroborating evidence \u2014 the height-diameter asymmetry "
        "in paired discrepancies, the 6:1 anomaly ratio in remeasurement "
        "data, and the concentration of estimation on the hard-to-verify "
        "measurement dimension \u2014 aligns with the multi-task "
        "prediction about where effort reduction concentrates. The "
        "allometric conformity null indicates that estimation does not "
        "produce detectable forensic signatures through residual analysis, "
        "a finding with direct implications for monitoring system design "
        "(Section 4). Additional forensic analyses \u2014 including digit "
        "heaping, allometric residual patterns, a "
        "difference-in-differences comparison using QA crews as a control "
        "group, and remeasurement growth anomaly analysis across 2.5 "
        "million tree pairs \u2014 are reported in the Supplementary "
        "Materials.",

        "Table A1 in the Supplementary Materials systematically catalogs "
        "discriminating predictions across all analyses. Of twelve tests, "
        "the two strongest \u2014 the 4:1 HTCD differential and "
        "height-diameter discrepancy asymmetry \u2014 directly support "
        "differential estimation. The seasonal forensic tests largely "
        "return nulls, consistent with year-round estimation rather than "
        "late-season concentration. Two results \u2014 allometric "
        "conformity (p = 0.44) and late-season variance increase \u2014 "
        "are most difficult to reconcile with a simple estimation model, "
        "suggesting that estimation is performed with sufficient skill to "
        "be forensically undetectable at the population level. We "
        "encourage readers to consult Table A1 for the complete pattern "
        "of evidence.",
    ]
    for p in paras:
        doc.add_paragraph(p)


def add_section4(doc: Document):
    """Implications (renumbered from old Section 5)."""
    add_heading(doc, "4. Implications for Timber Markets and Data Governance")

    paras = [
        "If ocular estimation occurs as the model predicts, its effects on timber "
        "inventory estimates are not random. Approximated measurements are "
        "The paired data provide a nuanced picture of estimation\u2019s consequences: "
        "estimated heights produce allometric residuals statistically "
        "indistinguishable from measured heights (Section 3.3), indicating that "
        "the variance-compression mechanism is not detectable at the population "
        "level. To the extent that estimation reduces the biological information "
        "content of height measurements, its effects on volume estimates would be "
        "non-random: height approximation errors would understate volume for tall "
        "trees and overstate it for short trees, with the net aggregate bias "
        "potentially modest if deviations approximately cancel. The data quality "
        "concern extends beyond systematic volume bias to loss of biological "
        "variance information, attenuation of the HTCD method code signal, and "
        "institutional corrosion of the monitoring architecture.",

        "FIA data propagate through multiple downstream applications. Stumpage "
        "pricing uses local volume estimates derived from FIA data; to the extent "
        "that estimation reduces biological information content, appraisals "
        "may understate merchantable volume on sites where trees exceed allometric "
        "expectations. National forest carbon stock estimates rely on allometric "
        "biomass equations applied to FIA measurements; to the extent that "
        "estimation compresses variance around the allometric mean, it would "
        "reduce detected spatial heterogeneity. Forest health monitoring\u2019s "
        "ancillary measurements \u2014 crown condition, damage codes, mortality \u2014 "
        "are arguably more susceptible to approximation, as the multi-task "
        "framework (Section 2.5) predicts.",

        "The ocular estimation problem is not unique to forest inventory. Analogous "
        "structures arise in fisheries observer programs, environmental compliance "
        "monitoring (Duflo et al., 2013), agricultural crop surveys, and health "
        "surveys. What distinguishes the FIA setting is the completeness of "
        "institutional information available for calibration: federal pay "
        "schedules, documented QA structure, and precisely timed measurement data.",

        "More broadly, empirical studies that rely on inventory-derived volume "
        "estimates \u2014 including timber market analyses, carbon accounting models, "
        "and forest growth studies \u2014 could in principle be affected by "
        "systematic measurement error of the type characterized here. Critically, "
        "this error is non-classical: rather than adding noise symmetrically, "
        "ocular estimation compresses height values toward the allometric mean, "
        "attenuating biological variance in ways that standard measurement error "
        "corrections do not address. Characterizing these downstream effects "
        "across specific applications is a natural direction for future work, "
        "though the magnitude in any given study remains to be determined.",
    ]
    for p in paras:
        doc.add_paragraph(p)


def add_section5(doc: Document):
    """Policy Recommendations (renumbered from old Section 6)."""
    add_heading(doc, "5. Policy Recommendations")

    paras = [
        "The calibrated model provides a framework for evaluating specific "
        "interventions. These recommendations flow from the institutional "
        "analysis, model calibration, and the empirical finding that production "
        "crews estimate at approximately four times the QA crew rate.",
    ]
    for p in paras:
        doc.add_paragraph(p)

    add_heading(doc, "Wage Adjustments", level=3)
    doc.add_paragraph(
        "At central parameters, the NSC critical wage is approximately $30.69/hr "
        "\u2014 far above current GS-4 ($17.27) and GS-6 ($21.54) wages. Simply "
        "raising wages to satisfy the NSC at current detection probabilities is "
        "infeasible within the General Schedule. However, at q = 0.04, the "
        "critical wage falls to $22.91/hr, within reach of GS-7 ($23.94/hr). "
        "Increasing effective monitoring is more cost-effective than raising wages."
    )

    add_heading(doc, "Increasing Detection Probability", level=3)
    paras = [
        "Expanded blind-check QA. Increasing the audit rate from 4% to 10% would "
        "approximately double the effective per-period detection probability, "
        "reducing the efficiency wage premium by roughly half. The national FIA "
        "plot network consists of approximately 130,000 forested plots. Increasing "
        "from 4% to 10% requires approximately 7,800 additional QA plot visits "
        "(130,000 \u00d7 0.06). At approximately 17 QA plots per crew-week, this "
        "requires roughly 450 additional crew-weeks, at approximately $4,000 "
        "per crew-week (two GS-7 technicians at $23.94/hr plus travel and per "
        "diem), or approximately $1.8 million/year nationally.",

        "Technology-assisted monitoring. GPS track logging and accelerometer data "
        "from field devices can verify plot visits and time on-plot. The marginal "
        "cost is near zero if crews already carry GPS-enabled devices.",

        "Automated consistency checks. Machine learning models trained on verified "
        "measurements can flag anomalous submissions in near-real-time \u2014 for "
        "instance, height-diameter relationships that are implausibly tight "
        "(suggesting estimation), increasing the subjective detection probability.",
    ]
    for p in paras:
        doc.add_paragraph(p)

    add_heading(doc, "Institutional Redesign", level=3)
    paras = [
        "Extended appointments with rehire guarantees would reduce the "
        "end-of-season separation rate b and increase the value of continued "
        "employment. Performance feedback loops linking QA results to individual "
        "records would increase both effective detection probability and "
        "consequences of detection. Career ladders from GS-4/5 seasonal to "
        "GS-7/9 permanent positions would align intrinsic motivation with "
        "institutional incentives.",

        "Finally, remote sensing (e.g., LiDAR-derived tree heights) may "
        "eventually eliminate the human measurement effort margin for the "
        "variable most susceptible to approximation, though ground-based "
        "measurement remains essential for species identification, crown "
        "condition, and other variables that remote sensing cannot reliably "
        "estimate.",
    ]
    for p in paras:
        doc.add_paragraph(p)

    add_heading(doc, "Feasibility Under Current Federal Constraints", level=3)
    paras = [
        "These recommendations vary substantially in feasibility under current "
        "federal workforce conditions. We sequence them by implementation "
        "difficulty: (1) Technology-assisted monitoring (near-zero marginal cost) "
        "\u2014 most robust to any budget scenario; (2) HTCD-based auditing "
        "(low cost) \u2014 automatable within existing QA infrastructure; "
        "(3) Expanded blind-check QA (~$1.8M/yr) \u2014 moderate cost, requires "
        "sustained appropriations; (4) Career ladders and extended appointments "
        "(highest cost) \u2014 addresses root cause but most vulnerable to hiring "
        "freezes.",

        "A potential vicious cycle deserves note: workforce instability "
        "simultaneously causes the data quality problem (by creating conditions "
        "that violate the NSC) and prevents its solution (by constraining the "
        "agency\u2019s capacity to implement reforms). Breaking this cycle requires "
        "recognizing that investment in monitoring infrastructure protects the "
        "downstream value of the data asset.",
    ]
    for p in paras:
        doc.add_paragraph(p)


def add_section6(doc: Document):
    """Limitations (renumbered from old Section 7)."""
    add_heading(doc, "6. Limitations")

    paras = [
        "Geographic and temporal scope. The paired QA/production analysis covers "
        "24 northern states during 2011\u20132016. Whether the HTCD differential "
        "varies by region remains an open question, though consistency across "
        "24 states suggests a structural pattern. The HTCD field records the "
        "crew\u2019s self-reported method code; if crews systematically underreport "
        "estimation, the 23.1% production rate is a lower bound.",

        "Selection versus incentives. QA and production crews may differ on "
        "dimensions beyond monitoring incentives \u2014 experience, training "
        "intensity, supervisory oversight. The HTCD differential reflects "
        "institutional differences broadly, not incentive effects in isolation. "
        "The policy prescription holds under multiple mechanisms, but the "
        "precise causal channel is not identified.",

        "Authorized estimation fraction unknown. FIA protocols permit visual "
        "estimation in specific field conditions. The differential relative to "
        "QA crews provides a lower bound on excess estimation, but the absolute "
        "rate of unauthorized estimation is not identified.",

        "Binary effort assumption. The binary effort choice overstates "
        "discreteness. The continuous-effort extension (Holmstr\u00f6m, 1979) "
        "yields qualitatively similar predictions. Our calibration results "
        "indicate directional incentive pressures rather than precise wage "
        "thresholds.",

        "Effort cost parameterization. The effort cost \u0113 is grounded in a "
        "time-budget micro-foundation (\u0113/w \u2208 [0.13, 0.36]), but the "
        "wide range reflects substantial terrain and canopy variation. "
        "Of the calibrated parameters, the effort cost has the weakest direct "
        "empirical grounding. Specific NSC percentage gaps (e.g., the 78% and "
        "43% benchmarks) should be interpreted as indicative of direction and "
        "approximate magnitude rather than precise thresholds. Qualitative "
        "conclusions are robust across this range (Table 2).",

        "Crew composition and self-selection. We lack crew-level identifiers "
        "and cannot distinguish within-crew from between-crew composition "
        "effects. Career-track seasonals\u2019 tendency to stay late means the "
        "seasonal signal predicted by the model would be attenuated even if "
        "terminal seasonals engage in estimation.",

        "Generalizability. Our analysis is specific to the U.S. FIA program. "
        "Other national forest inventories may face different incentive "
        "structures depending on workforce model, monitoring intensity, and "
        "labor markets.",
    ]
    for p in paras:
        doc.add_paragraph(p)


def add_section7(doc: Document):
    """Conclusion (renumbered from old Section 8)."""
    add_heading(doc, "7. Conclusion")

    paras = [
        "This paper documents differential measurement effort in the Forest "
        "Inventory and Analysis program using paired QA/production data: "
        "production crews record ocular estimation for 23.1% of heights versus "
        "5.8% for QA crews \u2014 a 4:1 ratio that is pervasive year-round. "
        "Height discrepancies between crew types are dramatic (3.31 ft) while "
        "diameter discrepancies are minimal (0.055 in.), confirming that effort "
        "reduction concentrates on the hard-to-verify measurement margin as "
        "multi-task theory predicts. A supporting efficiency wage calibration "
        "shows that current technician wages fall substantially below the "
        "no-shirking condition under central parameter assumptions "
        "(approximately 78% at GS-4 and 43% at GS-6, though sensitive to "
        "the effort cost parameter), with the detection probability identified "
        "as the dominant parameter.",

        "Three findings deserve emphasis. First, differential estimation is a "
        "structural policy design failure, not a moral failure of individual "
        "technicians. When detection probabilities are low, seasonal employment "
        "offers no career stake, and outside options are competitive, any rational "
        "agent faces incentives to reduce costly measurement effort. "
        "As Section 3 documents, workload pressure \u2014 the roughly 3:1 time-cost "
        "ratio between clinometer height measurement and DBH taping \u2014 provides "
        "an equally viable mechanism; crucially, both the incentive and workload "
        "explanations point toward the same institutional redesign response.",

        "Second, increasing effective detection probability is more cost-effective "
        "than raising wages. Increasing the per-period detection probability from "
        "0.01 to 0.04 would reduce the NSC critical wage from approximately "
        "$31/hr to $23/hr under the benchmark calibration, within reach of "
        "current GS-7 pay grades ($23.94/hr).",

        "Third, systematic auditing of HTCD method code distributions across "
        "crews and seasons could identify units with anomalously high estimation "
        "rates, enabling targeted intervention at minimal cost.",

        "The ocular estimation problem extends beyond forest inventory. Any "
        "monitoring program relying on seasonal field workers under imperfect "
        "observation faces analogous incentive structures. The broader "
        "implication is that institutional design matters: monitoring programs "
        "that rely primarily on intrinsic motivation without aligning extrinsic "
        "incentives are structurally vulnerable to data quality degradation. "
        "While our analysis is specific to the U.S. FIA program, the "
        "institutional structures that create incentives for measurement "
        "approximation \u2014 seasonal workforces, imperfect monitoring, "
        "hard-to-verify outputs \u2014 are common across public data collection "
        "programs worldwide, suggesting that the analytical framework developed "
        "here has broad applicability to monitoring program design.",
    ]
    for p in paras:
        doc.add_paragraph(p)


def add_references(doc: Document):
    """References section."""
    add_heading(doc, "References")

    refs = [
        "Akerlof, George A., and Janet L. Yellen, eds. 1986. Efficiency Wage Models of the Labor Market. Cambridge University Press.",
        "Bechtold, William A., and Paul L. Patterson, eds. 2005. The Enhanced Forest Inventory and Analysis Program \u2014 National Sampling Design and Estimation Procedures. USDA Forest Service GTR SRS-80.",
        "B\u00e9nabou, Roland, and Jean Tirole. 2003. \u201cIntrinsic and Extrinsic Motivation.\u201d Review of Economic Studies 70(3): 489\u2013520.",
        "B\u00e9nabou, Roland, and Jean Tirole. 2006. \u201cIncentives and Prosocial Behavior.\u201d American Economic Review 96(5): 1652\u20131678.",
        "Buurman, Margaretha, Josse Delfgaauw, Robert Dur, and Seth Van den Bossche. 2012. \u201cPublic Sector Employees: Risk Averse and Altruistic?\u201d Journal of Economic Behavior & Organization 83(3): 279\u2013291.",
        "Duflo, Esther, Michael Greenstone, Rohini Pande, and Nicholas Ryan. 2013. \u201cTruth-telling by Third-party Auditors and the Response of Polluting Firms: Experimental Evidence from India.\u201d Quarterly Journal of Economics 128(4): 1499\u20131545.",
        "Frey, Bruno S., and Reto Jegen. 2001. \u201cMotivation Crowding Theory.\u201d Journal of Economic Surveys 15(5): 589\u2013611.",
        "Holmstr\u00f6m, Bengt. 1979. \u201cMoral Hazard and Observability.\u201d Bell Journal of Economics 10(1): 74\u201391.",
        "Holmstr\u00f6m, Bengt, and Paul Milgrom. 1991. \u201cMultitask Principal-Agent Analyses.\u201d Journal of Law, Economics, & Organization 7(Special Issue): 24\u201352.",
        "Mosimann, James E., Claire V. Wiseman, and Ruth E. Eddy. 2002. \u201cData Fabrication: Can People Generate Random Digits?\u201d Accountability in Research 9(1): 31\u201345.",
        "Nigrini, Mark J. 1996. \u201cA Taxpayer Compliance Application of Benford\u2019s Law.\u201d Journal of the American Taxation Association 18(1): 72\u201391.",
        "Olken, Benjamin A. 2007. \u201cMonitoring Corruption: Evidence from a Field Experiment in Indonesia.\u201d Journal of Political Economy 115(2): 200\u2013249.",
        "Perry, James L., and Lois Recascino Wise. 1990. \u201cThe Motivational Bases of Public Service.\u201d Public Administration Review 50(3): 367\u2013373.",
        "Shapiro, Carl, and Joseph E. Stiglitz. 1984. \u201cEquilibrium Unemployment as a Worker Discipline Device.\u201d American Economic Review 74(3): 433\u2013444.",
        "Simonsohn, Uri. 2013. \u201cJust Post It: The Lesson from Two Cases of Fabricated Data Detected by Statistics Alone.\u201d Psychological Science 24(10): 1875\u20131888.",
        "Tomppo, Erkki, Thomas Gschwantner, Mark Lawrence, and Ronald E. McRoberts, eds. 2010. National Forest Inventories: Pathways for Common Reporting. Springer.",
        "USDA Forest Service. 2023. FIA National Core Field Guide, Volume 1. Washington, DC.",
        "Yanai, Ruth D., Alexander R. Young, John L. Campbell, James A. Westfall, Charles J. Barnett, Gretchen A. Dillon, Mark B. Green, and Christopher W. Woodall. 2023. \u201cMeasurement Uncertainty in a National Forest Inventory: Results from the Northern Region of the USA.\u201d Canadian Journal of Forest Research 53(3): 163\u2013177.",
        "Yanai, Ruth D., et al. 2023. \u201cPaired measurements from quality assurance visits to Forest Inventory and Analysis plots in the northern United States.\u201d Fort Collins, CO: Forest Service Research Data Archive. doi:10.2737/RDS-2022-0056.",
    ]
    for ref in refs:
        p = doc.add_paragraph(ref)
        p.paragraph_format.first_line_indent = Inches(-0.5)
        p.paragraph_format.left_indent = Inches(0.5)


def add_appendix(doc: Document):
    """Supplementary Materials: Forensic Analyses."""
    doc.add_page_break()
    add_heading(doc, "Supplementary Materials: Forensic Analyses")

    doc.add_paragraph(
        "This appendix presents four forensic statistical tests using public FIA "
        "data, a summary table of discriminating predictions, and robustness "
        "analyses. These supplement the paired QA/production comparison in "
        "Section 3 of the main text."
    )

    add_heading(doc, "A.1 FIA DataMart Data", level=2)
    doc.add_paragraph(
        "We use tree-level data from the FIA DataMart for eight states spanning "
        "five FIA regions: Vermont and Maine (Northeast), Minnesota and Wisconsin "
        "(Northern), Georgia (South), Colorado (Rocky Mountain), and Oregon and "
        "Washington (Pacific Northwest). Restricting to production plots "
        "(QA_STATUS = 1) and live trees with valid DBH yields 4,628,494 "
        "tree-observations (1995\u20132025). For the QA comparison, we extract "
        "91,493 additional trees from QA blind-check plots (QA_STATUS = 7). "
        "Georgia contributes the largest QA share (55,345 trees). "
        "LATE_SEASON is defined as in Section 3.1 (indicator for months "
        "9\u201312). All regressions use OLS with standard errors clustered "
        "at the state-year level."
    )

    add_heading(doc, "A.2 Test 1: Digit Heaping", level=2)
    doc.add_paragraph(
        "If technicians approximate DBH, measurements cluster on round numbers. "
        "The overall whole-inch rate is 11.57%, significantly above the 10% "
        "uniform expectation (\u03c7\u00b2 = 41,283, p < 0.001). However, the "
        "seasonal pattern is weak: LATE_SEASON coefficient = 0.0005 "
        "(SE = 0.0003, p = 0.135). Height rounding is also not significant "
        "(p = 0.125). Heaping is concentrated at digits 0 and 1 near the "
        "5.0-inch merchantability threshold, suggesting rounding conventions "
        "rather than seasonal effort variation."
    )

    add_figure(doc, FIG_DIR / "digit_heaping.png",
               "Figure A1: Digit Heaping by Measurement Month. "
               "(a) Whole-inch DBH rate; (b) Height rounding rate.")
    add_figure(doc, FIG_DIR / "last_digit_distribution.png",
               "Figure A2: Distribution of DBH Last Digit by Season")

    add_heading(doc, "A.3 Test 2: Allometric Residual Patterns", level=2)
    doc.add_paragraph(
        "If technicians estimate heights from diameter, allometric residuals "
        "will have lower variance (estimates track the curve). We fit the model "
        "with species group FE on production data (N = 4,587,949, R\u00b2 = 0.848). "
        "Levene\u2019s test finds significantly lower late-season variance "
        "(F = 1,642, p < 0.001; SD = 0.219 vs. 0.228). However, the regression "
        "with state and year FE yields \u03b1\u2081 \u2248 0 (p = 0.93), suggesting "
        "compositional confounds."
    )

    add_figure(doc, FIG_DIR / "allometric_residuals.png",
               "Figure A3: Allometric Residual Patterns by Month. "
               "(a) Residual standard deviation; (b) Month-specific R\u00b2.")

    add_heading(doc, "A.4 Test 3: QA Crew Comparison (Difference-in-Differences)", level=2)
    doc.add_paragraph(
        "QA crews (QA_STATUS = 7) perform blind remeasurements under "
        "fundamentally different institutional conditions. We estimate: "
        "|\u03b5| = \u03b1\u2080 + \u03b1\u2081\u00b7QA + \u03b1\u2082\u00b7LATE "
        "+ \u03b1\u2083\u00b7(QA \u00d7 LATE) + state FE + year FE + \u03b7."
    )

    # DID Results Table
    doc.add_paragraph()
    caption = doc.add_paragraph()
    run = caption.add_run("Table A1: Difference-in-Differences Regression Results")
    run.bold = True
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER

    table = doc.add_table(rows=6, cols=7)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ["", "QA", "SE", "LATE", "SE", "QA \u00d7 LATE", "p"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True

    did_data = [
        ["(1) |Residual|", "\u22120.0057", "(0.0018)", "0.0001", "(0.0008)",
         "0.0077**", "0.026"],
        ["(2) Residual\u00b2", "\u22120.0025", "(0.0011)", "0.0001", "(0.0005)",
         "0.0039*", "0.056"],
        ["(3) |Resid| + Spp FE", "\u22120.0045", "(0.0016)", "\u22120.0004", "(0.0008)",
         "0.0067**", "0.045"],
        ["(4) DBH Heaping", "\u22120.0001", "(0.0011)", "0.0005", "(0.0003)",
         "0.0021", "0.376"],
        ["(5) HT Rounding", "0.0028", "(0.0026)", "\u22120.0019", "(0.0012)",
         "0.0052", "0.270"],
    ]
    for row_idx, row_data in enumerate(did_data, start=1):
        for col_idx, val in enumerate(row_data):
            table.rows[row_idx].cells[col_idx].text = val

    note = doc.add_paragraph()
    note.paragraph_format.space_before = Pt(2)
    run = note.add_run(
        "Notes: N = 4,679,419 (Models 1\u20133, 5) or 4,719,987 (Model 4). "
        "All models include state and year fixed effects. Standard errors "
        "clustered at state-year level. ** p < 0.05, * p < 0.10."
    )
    run.font.size = Pt(10)
    run.italic = True

    doc.add_paragraph(
        "The DID interaction is positive: \u03b1\u2083 = 0.0077 (SE = 0.0035, "
        "p = 0.026). However, the wild cluster bootstrap yields p = 0.403, "
        "indicating this result does not survive inference appropriate to "
        "8 effective clusters."
    )

    add_figure(doc, FIG_DIR / "qa_comparison_residuals.png",
               "Figure A4: Allometric Residual Dispersion by Month: "
               "Production vs. QA Crews. (a) Residual SD; (b) Mean |residual|.")

    add_heading(doc, "A.5 Test 4: Remeasurement Growth Anomalies", level=2)
    doc.add_paragraph(
        "Across 2,464,859 remeasurement pairs in eight states, the height "
        "anomaly rate is 9.3% compared to 1.5% for DBH \u2014 a 6:1 ratio. "
        "The conifer-only LATE_SEASON coefficient is effectively zero "
        "(p = 0.98). Late-season height growth variance is higher for both "
        "species groups, opposite to the ocular estimation prediction."
    )

    add_heading(doc, "A.6 Discriminating Predictions", level=2)
    doc.add_paragraph(
        "Table A2 summarizes the discriminating predictions and observed "
        "results across all analyses."
    )

    t_pred = doc.add_table(rows=9, cols=4)
    t_pred.style = "Table Grid"
    t_pred.alignment = WD_TABLE_ALIGNMENT.CENTER
    pred_headers = ["Test", "Ocular Est.", "Honest Error", "Observed"]
    for i, h in enumerate(pred_headers):
        t_pred.rows[0].cells[i].text = h
        for p in t_pred.rows[0].cells[i].paragraphs:
            for r in p.runs:
                r.bold = True

    pred_data = [
        ["Digit heaping increases late season",
         "Yes", "No", "Not sig. (p = 0.13)"],
        ["Resid. variance decreases late season",
         "Yes", "No (increases)", "Sig. unconditionally"],
        ["DID: Prod. suppresses variance increase",
         "Yes", "No", "Conv. p=0.026; bootstrap p=0.403"],
        ["HT anomaly rate > DBH (remeasurement)",
         "Yes", "Yes (HT harder)", "Yes \u2014 9.3% vs 1.5%"],
        ["Conifer HT anomalies increase late",
         "Yes", "No change", "No \u2014 coef \u2248 0, p = 0.98"],
        ["Paired: Prod. HTCD=3 > QA",
         "Yes", "No", "Yes \u2014 23.1% vs 5.8% (4:1)"],
        ["Paired: HT discrep. > DIA discrep.",
         "Yes", "Yes (HT harder)", "Yes \u2014 3.31 ft vs 0.06 in."],
        ["Paired: Est. HT smaller allom. resid.",
         "Yes (track curve)", "No", "No \u2014 identical (p = 0.44)"],
    ]
    for row_idx, row_data in enumerate(pred_data, start=1):
        for col_idx, val in enumerate(row_data):
            t_pred.rows[row_idx].cells[col_idx].text = val

    cap_pred = doc.add_paragraph()
    run = cap_pred.add_run("Table A2: Discriminating Predictions and Results")
    run.bold = True
    run.italic = True
    cap_pred.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_heading(doc, "A.7 Robustness", level=2)

    # Sensitivity table
    sens_headers = ["Sample", "QA Trees", "\u03b1\u2083 (QA \u00d7 LATE)", "SE", "p"]
    sens_rows = [
        ["Full sample (8 states)", "91,493", "0.0077", "0.0035", "0.026"],
        ["Georgia only", "55,345", "0.0053", "0.0044", "0.228"],
        ["CO + OR + WA", "32,709", "\u22120.0019", "0.0093", "0.838"],
        ["All except GA (7 states)", "36,125", "0.0027", "0.0079", "0.729"],
    ]
    t_sens = doc.add_table(rows=len(sens_rows) + 1, cols=len(sens_headers))
    t_sens.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(sens_headers):
        cell = t_sens.rows[0].cells[i]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True
    for row_idx, row in enumerate(sens_rows, 1):
        for col_idx, val in enumerate(row):
            t_sens.rows[row_idx].cells[col_idx].text = val

    cap_sens = doc.add_paragraph()
    run = cap_sens.add_run(
        "Table A3: DID Sensitivity by State Grouping. "
        "All models include state and year FE. "
        "Dependent variable: |allometric residual|."
    )
    run.bold = True
    run.italic = True
    cap_sens.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(
        "No individual subsample produces a significant interaction. The wild "
        "cluster bootstrap (Cameron et al., 2008) yields p = 0.403 for the "
        "primary DID specification, reflecting the small effective cluster count "
        "(8 states) and Georgia\u2019s dominance of the QA sample. The paired "
        "analysis in Section 3 addresses the GA-dominance concern with "
        "24-state coverage."
    )


def add_csv_table(doc: Document, csv_path: Path, caption: str):
    """Insert a CSV as a formatted Word table with caption."""
    doc.add_paragraph()
    cap = doc.add_paragraph()
    run = cap.add_run(caption)
    run.bold = True
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER

    with open(csv_path) as f:
        reader = csv.reader(f)
        rows = list(reader)

    if not rows:
        return

    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            table.rows[i].cells[j].text = val.strip()
            if i == 0:
                for p in table.rows[i].cells[j].paragraphs:
                    for r in p.runs:
                        r.bold = True

    doc.add_paragraph()


def add_figure(doc: Document, img_path: Path, caption: str):
    """Insert a figure image with caption."""
    if not img_path.exists():
        doc.add_paragraph(f"[FIGURE NOT FOUND: {img_path.name}]")
        return

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(img_path), width=Inches(5.5))

    cap = doc.add_paragraph()
    run = cap.add_run(caption)
    run.italic = True
    run.font.size = Pt(10)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()


def main():
    doc = Document()

    # Page setup
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    set_style(doc)

    add_title_page(doc)
    add_abstract(doc)
    add_section1(doc)
    add_section2(doc)
    add_section3(doc)
    add_section4(doc)
    add_section5(doc)
    add_section6(doc)
    add_section7(doc)
    add_references(doc)
    add_appendix(doc)

    doc.save(str(OUT_PATH))
    print(f"Saved: {OUT_PATH}")
    print(f"Size: {OUT_PATH.stat().st_size / 1024:.0f} KB")


if __name__ == "__main__":
    main()
